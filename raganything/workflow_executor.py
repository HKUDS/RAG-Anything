"""
工作流执行引擎 — 拓扑排序 + 6×真实RAG节点执行器 + WebSocket 状态推送
"""
import json
import time
import uuid
from pathlib import Path
from datetime import datetime
from collections import deque

RUNS_DIR = Path("./workflows/runs")
RUNS_DIR.mkdir(parents=True, exist_ok=True)


class ExecutionContext:
    """执行上下文 — 注入运行时依赖（LLM/Embedding/LightRAG实例等）"""

    def __init__(
        self,
        *,
        llm_model: str = "",
        llm_api_key: str = "",
        llm_base_url: str = "",
        embed_model: str = "",
        embed_dim: int = 1024,
        embed_api_key: str = "",
        embed_base_url: str = "",
        kb_instance=None,  # RAGAnything instance
        upload_dir: Path = Path("./uploads"),
        openai_complete_func=None,
        openai_embed_func=None,
    ):
        self.llm_model = llm_model
        self.llm_api_key = llm_api_key
        self.llm_base_url = llm_base_url
        self.embed_model = embed_model
        self.embed_dim = embed_dim
        self.embed_api_key = embed_api_key
        self.embed_base_url = embed_base_url
        self.kb_instance = kb_instance
        self.upload_dir = Path(upload_dir)
        self.openai_complete_func = openai_complete_func
        self.openai_embed_func = openai_embed_func


class NodeExecutor:
    """执行器基类"""
    node_type: str = ""

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        raise NotImplementedError


# ──────────────────────────────────────────────
# 辅助函数
# ──────────────────────────────────────────────

def _extract_text(inputs: dict, *keys: str) -> str:
    """从上游节点输出中提取文本，按优先级依次尝试多个字段"""
    for key in keys:
        val = inputs.get(key, "")
        if isinstance(val, list):
            # 处理检索结果 [{chunk, score}, ...] 格式
            if val and isinstance(val[0], dict):
                val = "\n".join(
                    f"[相关度: {d.get('score', 'N/A')}] {d.get('chunk', str(d))}"
                    for d in val
                )
            else:
                val = "\n".join(str(v) for v in val)
        if val and str(val).strip():
            return str(val)
    # 兜底：拼接所有字段值
    parts = []
    for k, v in inputs.items():
        if k in ("duration_ms", "count", "dims", "format", "error", "file_type", "size_bytes",
                  "search_mode", "chunk_count", "mode", "top_k"):
            continue
        if isinstance(v, list):
            if v and isinstance(v[0], dict):
                parts.extend(
                    f"[相关度: {d.get('score', 'N/A')}] {d.get('chunk', str(d))}"
                    for d in v
                )
            else:
                parts.extend(str(x) for x in v)
        elif isinstance(v, str) and v.strip():
            parts.append(v)
    return "\n".join(parts) if parts else ""


# ──────────────────────────────────────────────
# 6 种节点执行器（现已连接真实 RAG 组件）
# ──────────────────────────────────────────────

class DocumentInputExecutor(NodeExecutor):
    """文档输入 — 从 uploads/ 目录读取真实文件，支持指定文件名"""
    node_type = "document_input"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        file_type = config.get("file_type", ".pdf")
        file_name = config.get("file_name", "")
        max_size_mb = config.get("max_size_mb", 100)
        max_size = max_size_mb * 1024 * 1024

        upload_dir = ctx.upload_dir
        upload_dir.mkdir(exist_ok=True)

        # 指定文件 → 直接读；未指定 → 选最新匹配类型的文件
        if file_name:
            target = upload_dir / file_name
            if not target.exists():
                return {"content": "", "file_name": file_name, "error": f"文件不存在: {file_name}"}
        else:
            candidates = sorted(
                [f for f in upload_dir.iterdir() if f.is_file() and f.suffix.lower() == file_type],
                key=lambda p: p.stat().st_mtime, reverse=True,
            ) if file_type != "全部" else sorted(
                [f for f in upload_dir.iterdir() if f.is_file()],
                key=lambda p: p.stat().st_mtime, reverse=True,
            )
            if not candidates:
                return {"content": "", "file_name": "", "error": f"uploads/ 中没有 {file_type} 文件"}
            target = candidates[0]

        if target.stat().st_size > max_size:
            return {"content": "", "file_name": target.name, "error": f"文件超过 {max_size_mb}MB 限制"}

        try:
            if target.suffix.lower() in ('.txt', '.md', '.csv', '.json', '.py', '.js', '.html', '.xml'):
                text = target.read_text(encoding='utf-8', errors='replace')
            elif target.suffix.lower() == '.pdf':
                # 尝试用 pypdfium2 读取 PDF 文本
                try:
                    import pypdfium2 as pdfium
                    pdf = pdfium.PdfDocument(str(target))
                    pages = [pdf[i].get_textpage().get_text_range() for i in range(len(pdf))]
                    text = '\n\n'.join(pages)
                except Exception:
                    # fallback: 读取为二进制 + decode
                    text = f"[PDF] {target.name} (需 MinerU/Docling 解析器)"
            elif target.suffix.lower() == '.docx':
                try:
                    from docx import Document
                    doc = Document(str(target))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    text = '\n'.join(paragraphs)
                except Exception:
                    text = f"[DOCX] {target.name}\n路径: {target.absolute()}"
            else:
                # .docx, .xlsx 等 — 返回文件路径供后续解析器处理
                text = f"[{target.suffix}] {target.name}\n路径: {target.absolute()}"
            return {"content": text, "file_name": target.name, "file_type": target.suffix, "size_bytes": target.stat().st_size}
        except Exception as e:
            return {"content": "", "file_name": target.name, "error": str(e)}


class TextSplitterExecutor(NodeExecutor):
    """文本分割 — 真实分块"""
    node_type = "text_splitter"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        text = _extract_text(inputs, "content", "chunks")
        if not text.strip():
            return {"chunks": [], "count": 0, "error": "上游节点没有提供文本内容"}

        chunk_size = int(config.get("chunk_size", 800))
        chunk_overlap = int(config.get("chunk_overlap", 100))
        step = max(1, chunk_size - chunk_overlap)
        chunks = []
        for i in range(0, len(text), step):
            chunk = text[i:i + chunk_size]
            if len(chunk.strip()) >= 20:
                chunks.append(chunk)
        if not chunks:
            chunks = [text[:chunk_size]]
        return {"chunks": chunks, "count": len(chunks)}


class EmbeddingExecutor(NodeExecutor):
    """嵌入向量 — 调用真实 embedding API"""
    node_type = "embedding"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        model = config.get("model") or ctx.embed_model or "text-embedding-v4"
        text = _extract_text(inputs, "content", "chunks", "results")
        if not text.strip():
            return {"vector": [], "error": "上游节点没有提供文本"}

        if ctx.openai_embed_func:
            try:
                # 使用 .func 绕过 @wrap_embedding_func_with_attrs 装饰器，
                # 否则会自动注入 embedding_dim=1536，与 DashScope text-embedding-v4 (1024) 冲突
                embed_fn = getattr(ctx.openai_embed_func, 'func', ctx.openai_embed_func)
                result = await embed_fn(
                    [str(text)[:8000]],  # 截断，避免超 token 限制
                    model=model,
                    api_key=ctx.embed_api_key,
                    base_url=ctx.embed_base_url,
                )
                # result 可能是 list 或 numpy 数组，不能直接用 if result 判空
                if result is not None and len(result) > 0:
                    vec = result[0]
                    if hasattr(vec, 'tolist'):
                        vec = vec.tolist()
                    return {"vector": vec, "model": model, "dims": len(vec)}
                return {"vector": [], "model": model, "dims": 0}
            except Exception as e:
                return {"vector": [], "model": model, "error": f"Embedding 调用失败: {e}"}
        return {"vector": [], "model": model, "error": "Embedding 函数未配置"}


def _cosine_similarity_top_k(query_vector: list, chunk_vectors: list, chunks: list, top_k: int) -> list[dict]:
    """用 numpy 做余弦相似度 Top-K 检索，返回 [{chunk, score}, ...]"""
    import numpy as np

    q = np.array(query_vector, dtype=np.float32).flatten()
    q_norm = q / (np.linalg.norm(q) + 1e-10)

    c = np.array(chunk_vectors, dtype=np.float32)
    if c.ndim == 1:
        c = c.reshape(1, -1)
    c_norm = c / (np.linalg.norm(c, axis=1, keepdims=True) + 1e-10)

    scores = np.dot(c_norm, q_norm)
    top_indices = np.argsort(scores)[::-1][:top_k]

    results = []
    for idx in top_indices:
        score = float(scores[idx])
        if score > 0 and idx < len(chunks):
            results.append({"chunk": str(chunks[idx])[:500], "score": round(score, 4)})
    return results


class RetrieverExecutor(NodeExecutor):
    """检索器 — 优先内存向量检索 (有上游 chunks)，降级 LightRAG 知识库"""
    node_type = "retriever"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        top_k = int(config.get("top_k", 10))
        mode = config.get("mode", "hybrid")

        # query_text 手动输入优先，否则从上游提取
        query = (config.get("query_text") or "").strip()
        if not query:
            query = _extract_text(inputs, "content", "chunks", "results")
        query = str(query)[:1000]

        # 检查上游是否提供了 chunks（内存检索模式）
        chunks = inputs.get("chunks")
        has_chunks = isinstance(chunks, list) and len(chunks) > 0

        if has_chunks and ctx.openai_embed_func:
            return await self._in_memory_search(query, chunks, top_k, config, ctx)

        # 降级到 LightRAG 知识库
        return await self._kb_search(query, top_k, mode, ctx)

    async def _in_memory_search(self, query: str, chunks: list, top_k: int, config: dict, ctx: ExecutionContext) -> dict:
        """内存向量检索：向量化 query+chunks（分批≤10），余弦相似度 Top-K"""
        import numpy as np
        embed_fn = getattr(ctx.openai_embed_func, 'func', ctx.openai_embed_func)
        embed_model = config.get("model") or ctx.embed_model or "text-embedding-v4"

        # 限制 chunk 数量，避免超 token
        max_chunks = min(len(chunks), 100)
        selected_chunks = chunks[:max_chunks]

        try:
            # 先向量化 query
            q_result = await embed_fn(
                [query],
                model=embed_model,
                api_key=ctx.embed_api_key,
                base_url=ctx.embed_base_url,
            )
            if q_result is None or len(q_result) == 0:
                return {"results": [], "query": query[:200], "search_mode": "in_memory",
                        "error": "Query 向量化失败"}
            query_vec = q_result[0]
            if hasattr(query_vec, 'tolist'):
                query_vec = query_vec.tolist()

            # 分批向量化 chunks（DashScope 限制 ≤10/批）
            BATCH = 10
            chunk_vecs = []
            for i in range(0, len(selected_chunks), BATCH):
                batch = selected_chunks[i:i + BATCH]
                c_result = await embed_fn(
                    batch,
                    model=embed_model,
                    api_key=ctx.embed_api_key,
                    base_url=ctx.embed_base_url,
                )
                if c_result is not None and len(c_result) > 0:
                    for v in c_result:
                        if hasattr(v, 'tolist'):
                            v = v.tolist()
                        chunk_vecs.append(v)

            if not chunk_vecs:
                return {"results": [], "query": query[:200], "search_mode": "in_memory",
                        "error": "Chunk 向量化失败"}

            search_results = _cosine_similarity_top_k(query_vec, chunk_vecs, selected_chunks, top_k)

            return {
                "results": search_results,
                "query": query[:200],
                "top_k": top_k,
                "search_mode": "in_memory",
                "chunk_count": len(chunks),
            }
        except Exception as e:
            return {"results": [], "query": query[:200], "search_mode": "in_memory",
                    "error": f"内存检索失败: {e}"}

    async def _kb_search(self, query: str, top_k: int, mode: str, ctx: ExecutionContext) -> dict:
        """LightRAG 知识库检索（降级路径）"""
        if not ctx.kb_instance or not hasattr(ctx.kb_instance, 'lightrag') or not ctx.kb_instance.lightrag:
            return {"results": [], "error": "知识库未初始化，请先上传文档到知识库",
                    "search_mode": "knowledge_base"}

        try:
            from lightrag import QueryParam
            result = await ctx.kb_instance.lightrag.aquery(
                query,
                param=QueryParam(mode=mode, top_k=top_k, only_need_context=True),
            )
            return {
                "results": [result] if isinstance(result, str) else result,
                "query": query[:200],
                "mode": mode,
                "top_k": top_k,
                "search_mode": "knowledge_base",
            }
        except Exception as e:
            return {"results": [], "error": f"检索失败: {e}", "query": query[:200],
                    "search_mode": "knowledge_base"}


class LLMAnswerExecutor(NodeExecutor):
    """LLM 回答 — 调用真实 LLM API"""
    node_type = "llm_answer"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        model = config.get("model") or ctx.llm_model or ""
        temperature = float(config.get("temperature", 0.1))
        system_prompt = config.get("system_prompt", "")

        # 提取问题（来自上游 retriever 的 query 字段）
        query = inputs.get("query", "").strip()
        # 提取上下文（检索结果）
        context = _extract_text(inputs, "content", "results", "chunks", "answer")
        if not context.strip():
            context = "（无可用上下文，请告知用户检查检索是否成功）"

        # 构建完整 prompt：问题 + 上下文
        prompt_parts = []
        if query:
            prompt_parts.append(f"问题：{query}")
        prompt_parts.append(f"基于以下上下文回答问题。如果上下文不足以回答问题，请如实说明。\n\n{str(context)[:8000]}")
        user_prompt = "\n\n".join(prompt_parts)

        if ctx.openai_complete_func:
            try:
                answer = await ctx.openai_complete_func(
                    model, user_prompt,
                    system_prompt=system_prompt or None,
                    api_key=ctx.llm_api_key,
                    base_url=ctx.llm_base_url,
                    temperature=temperature,
                    max_tokens=4096,
                )
                return {"answer": answer, "model": model, "temperature": temperature}
            except Exception as e:
                return {"answer": "", "error": f"LLM 调用失败: {e}", "model": model}
        return {"answer": "", "error": "LLM 函数未配置", "model": model}


class OutputExecutor(NodeExecutor):
    """输出 — 格式化上游结果"""
    node_type = "output"

    async def execute(self, config: dict, inputs: dict, ctx: ExecutionContext) -> dict:
        fmt = config.get("format", "markdown")
        content = _extract_text(inputs, "answer", "formatted", "content", "results", "chunks")
        if isinstance(content, list):
            content = "\n".join(f"- {c}" for c in content)
        if fmt == "markdown":
            formatted = f"## 工作流输出\n\n{content}"
        elif fmt == "json":
            formatted = f"```json\n{json.dumps({'output': str(content)}, ensure_ascii=False, indent=2)}\n```"
        else:
            formatted = str(content)
        return {"formatted": formatted, "format": fmt}


# ──────────────────────────────────────────────
# 注册表
# ──────────────────────────────────────────────

EXECUTORS: dict[str, NodeExecutor] = {
    "document_input": DocumentInputExecutor(),
    "text_splitter": TextSplitterExecutor(),
    "embedding": EmbeddingExecutor(),
    "retriever": RetrieverExecutor(),
    "llm_answer": LLMAnswerExecutor(),
    "output": OutputExecutor(),
}


# ──────────────────────────────────────────────
# 拓扑排序 + 工作流执行
# ──────────────────────────────────────────────

def topological_sort(nodes: list[dict], edges: list[dict]) -> list[str]:
    """拓扑排序，返回节点ID执行顺序。检测环路。"""
    node_ids = {n["id"] for n in nodes}
    in_degree = {nid: 0 for nid in node_ids}
    adj = {nid: [] for nid in node_ids}
    for e in edges:
        src, tgt = e["source"], e["target"]
        if src not in node_ids or tgt not in node_ids:
            continue
        adj[src].append(tgt)
        in_degree[tgt] += 1
    q = deque([nid for nid, deg in in_degree.items() if deg == 0])
    order = []
    while q:
        cur = q.popleft()
        order.append(cur)
        for nxt in adj[cur]:
            in_degree[nxt] -= 1
            if in_degree[nxt] == 0:
                q.append(nxt)
    if len(order) != len(node_ids):
        raise ValueError("Workflow contains a cycle")
    return order


async def execute_workflow(
    workflow: dict,
    *,
    ctx: ExecutionContext = None,
    status_callback=None,
) -> dict:
    """执行工作流 DAG"""
    nodes = workflow.get("nodes", [])
    edges = workflow.get("edges", [])
    node_map = {n["id"]: n for n in nodes}
    node_order = topological_sort(nodes, edges)

    run_id = str(uuid.uuid4())
    started_at = datetime.now().isoformat()
    node_results = []
    node_outputs = {}
    final_status = "completed"

    if ctx is None:
        ctx = ExecutionContext()
    ctx._run_id_cache = run_id  # 供外部 status_callback 使用

    async def push_status(node_id, status, data=None):
        node_results.append({"node_id": node_id, "status": status, "data": data or {}, "timestamp": datetime.now().isoformat()})
        if status_callback:
            await status_callback(node_id, status, data)

    for node_id in node_order:
        node = node_map[node_id]
        node_type = node.get("data", {}).get("nodeType", "")
        config = node.get("data", {})

        executor = EXECUTORS.get(node_type)
        if not executor:
            await push_status(node_id, "error", {"error": f"Unknown node type: {node_type}"})
            final_status = "failed"
            break

        # 收集上游节点输出
        upstream_inputs = {}
        for e in edges:
            if e["target"] == node_id:
                src_output = node_outputs.get(e["source"], {})
                upstream_inputs.update(src_output)

        await push_status(node_id, "running")
        start = time.time()
        try:
            result = await executor.execute(config, upstream_inputs, ctx)
            result["duration_ms"] = round((time.time() - start) * 1000)
            node_outputs[node_id] = result
            await push_status(node_id, "done", result)
        except Exception as exc:
            await push_status(node_id, "error", {"error": str(exc), "duration_ms": round((time.time() - start) * 1000)})
            final_status = "failed"
            break

    # 收集最终输出
    final_output = ""
    for node_id in reversed(node_order):
        output = node_outputs.get(node_id, {})
        if node_map[node_id].get("data", {}).get("nodeType") == "output":
            final_output = output.get("formatted", json.dumps(output, ensure_ascii=False))
            break
    if not final_output:
        last_out = node_outputs.get(node_order[-1], {}) if node_order else {}
        final_output = json.dumps(last_out, ensure_ascii=False, indent=2)

    completed_at = datetime.now().isoformat()
    run_record = {
        "run_id": run_id, "workflow_id": workflow.get("id", "unknown"),
        "workflow_name": workflow.get("name", "unnamed"),
        "status": final_status, "started_at": started_at, "completed_at": completed_at,
        "node_results": node_results, "final_output": final_output,
    }
    (RUNS_DIR / f"{run_id}.json").write_text(json.dumps(run_record, ensure_ascii=False, indent=2), encoding="utf-8")
    if status_callback:
        await status_callback(None, "run_complete", {"run_id": run_id, "status": final_status})
    return run_record
