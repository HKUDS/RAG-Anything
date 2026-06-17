#!/usr/bin/env python3
"""Generate 待开发功能详解 DOCX using python-docx — avoids JS quote encoding issues."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

for lvl, sz in [(1, 22), (2, 16), (3, 13)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = '微软雅黑'
    hs.font.size = Pt(sz)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    doc.add_paragraph(text)

def feature_table(rows, left_w=Cm(7.5)):
    """Two-column feature table: 功能描述 | 技术方案"""
    right_w = Cm(17) - left_w
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    # Header
    for ci, txt in enumerate(['功能描述', '技术方案与关键参数']):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        run = cell.paragraphs[0].add_run(txt)
        run.font.name = '微软雅黑'
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
    # Data rows
    for ri, (left_text, right_lines) in enumerate(rows):
        # Left cell
        lc = table.rows[ri + 1].cells[0]
        lc.text = ''
        run = lc.paragraphs[0].add_run(left_text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(9.5)
        run.bold = True
        if ri % 2 == 0:
            set_cell_shading(lc, 'F7F8FA')
        # Right cell
        rc = table.rows[ri + 1].cells[1]
        rc.text = ''
        for li, line in enumerate(right_lines):
            para = rc.paragraphs[0] if li == 0 else rc.add_paragraph()
            run = para.add_run(line)
            run.font.name = '微软雅黑'
            run.font.size = Pt(9.5)
        if ri % 2 == 0:
            set_cell_shading(rc, 'F7F8FA')
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════
# COVER
for _ in range(12):
    doc.add_paragraph()
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.add_run('RAG-Anything').font.size = Pt(28)
doc.add_paragraph()
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cp.add_run('待开发功能详解')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
doc.add_paragraph()
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.add_run('多模态知识库智能问答系统 · 后续版本规划').font.size = Pt(12)
doc.add_paragraph()
cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.add_run('版本 1.0 · 2026年6月').font.size = Pt(9)
doc.add_page_break()

# ═══ 一、开发总览 ═══
h('一、开发总览')
p('本文档为 RAG-Anything 后续版本的功能开发规格说明，涵盖第 2 周至第 4 周的交付项及远期规划。每项功能均标注可行性评估，并附技术方案、实现步骤和验收标准。')
p('当前已完成功能共计 16 项，详见《RAG-Anything 已实现功能说明书》。')

feature_table([
    ('第 2 周（进行中）', ['交付 3 项：HyDE + Multi-Query 查询改写模块、后端模块化重构（5 Router）、Agentic RAG 多步推理引擎。时间：2026.6.23 — 6.29。']),
    ('第 3 周（计划交付）', ['交付 6 项：RRF 显式三路融合检索、GraphRAG 知识图谱增强、DAG 工作流引擎（MVP 版）、SSO/OIDC 企业认证、多轮对话上下文管理、密钥管理与依赖安全扫描、核心测试覆盖。时间：2026.6.30 — 7.06。']),
    ('第 4 周（计划交付）', ['交付 5 项：审计日志系统、Docling/Marker 文档解析升级、前端 Zustand + i18n 重构、性能压测与生产调优、项目管理看板集成。时间：2026.7.07 — 7.13。']),
    ('远期规划', ['7 项：数据集模块、多数据库接入扩展、知识质量检测、非结构化数据打标应用、更多数据源接入、实时协作、移动端适配。时间待定，按需启动。']),
])

doc.add_page_break()

# ═══ 二、第 2 周 ═══
h('二、第 2 周交付项（进行中）')

h('2.1  HyDE + Multi-Query 查询改写独立模块', 2)
p('✅ 可行性：可行——本质是 LLM 文本生成 + 去重融合，基础改写代码已在 SSE 管线中存在。')
feature_table([
    ('功能目标', ['将当前内联在 SSE 流式管线中的查询改写功能封装为独立的可配置模块，新增 HyDE（假设文档嵌入）和 Multi-Query（多路变体生成）两种增强策略。']),
    ('HyDE 策略', ['对用户问题，先生成 3 个假设答案文本（temperature=0.7，每篇 max_tokens=512），再用这些假设答案的向量去检索。这种方式能有效弥补问题表述和答案表述之间的语义鸿沟。', '参数：n=3, temperature=0.7, max_tokens=512。']),
    ('Multi-Query 策略', ['将用户问题改写为 3 个不同表述的变体（temperature=0.8，每个 max_tokens=128），与原始问题一并检索，综合所有结果。', '参数：n=3, temperature=0.8, max_tokens=128。']),
    ('融合方式', ['原始查询 + 3 个 HyDE 假设文档 + 3 个 Multi-Query 变体 → 共 7 路并行检索 → 去重 → RRF 融合排序。单路超时 30 秒自动跳过。']),
    ('文件规划', ['新建 raganything/query_rewriter.py：QueryRewriter 类，含 rewrite_hyde()、rewrite_multiquery()、rewrite() 方法。在 server.py 的查询管线中集成调用。']),
    ('环境变量', ['QUERY_REWRITE_STRATEGIES=hyde,multiquery（控制启用策略）；QUERY_REWRITE_TIMEOUT=30（单策略超时秒数）。']),
    ('验收标准', ['输入年假政策→生成 3 个假设文档 + 3 个变体查询；召回率相比单查询提升 ≥ 20%；可通过环境变量关闭/启用各策略；超时自动降级为原始查询。']),
])

h('2.2  后端模块化重构（5 Router）', 2)
p('✅ 可行性：可行——纯代码搬运，不改变逻辑，每个 Router 文件控制在 400 行以内。')
feature_table([
    ('功能目标', ['将当前 server.py（约 2300 行单文件）拆分为 5 个独立 Router 文件，提升代码可维护性。']),
    ('Router 拆分方案', ['routers/auth.py（~300 行）：/api/auth/*；routers/knowledge.py（~400 行）：/api/knowledge/* + /api/upload/*；routers/agent.py（~350 行）：/api/agents/*；routers/query.py（~300 行）：/api/query/*；routers/admin.py（~300 行）：/api/admin/* + /api/settings + /api/monitor/*。']),
    ('公共依赖', ['将 limiter、get_current_user、verify_kb_access 等提取到 dependencies.py，各 Router 统一引用。']),
    ('验收标准', ['每个 Router 文件 < 400 行；server.py < 300 行；所有现有 API 端点路径不变；pytest 全部通过。']),
])

h('2.3  Agentic RAG 多步推理引擎', 2)
p('⚠️ 可行性：时间偏紧——完整 ReAct 循环需要 1.5-2 周。第 2 周先交付最小版本（2 步推理 + Search 工具），第 3 周完善。')
feature_table([
    ('功能目标', ['在现有智能体框架基础上实现 ReAct（Thought→Action→Observation 循环）多步推理模式。']),
    ('ReAct 推理循环', ['用户提问 → Thought（分析需要什么信息）→ Action（执行检索/计算）→ Observation（观察结果）→ 判断是否需要补充检索 → 循环直到信息充分或达到最大步数。参数：max_steps=5，单步超时 30 秒。']),
    ('内置工具', ['第 2 周：SearchTool（知识库检索）+ CalculatorTool（四则运算）。第 3 周追加：DatabaseQueryTool、WebSearchTool。工具抽象为 Tool 基类，支持注册自定义工具。']),
    ('文件规划', ['新建 raganything/agentic_rag.py：AgenticRAG 类 + Tool 基类。在查询端点中添加 agent_mode 参数。']),
    ('验收标准', ['MVP 版：Agent 分 2 步检索+计算；max_steps=5 时不会无限循环；单工具超时 30 秒后自动跳过。']),
])

doc.add_page_break()
# ═══ 三、第 3 周 ═══
h('三、第 3 周交付项')

h('3.1  RRF 显式三路融合检索', 2)
p('✅ 可行性：可行——RRF 公式仅一行代码，BM25 和向量搜索已实现，图谱通道基础已有。')
feature_table([
    ('功能目标', ['在现有 LightRAG hybrid 模式基础上，实现显式的三路独立检索通道 + RRF 融合排序。']),
    ('三路通道', ['BM25 关键词通道（权重 0.3）：k1=1.5, b=0.75, top_k=50；向量语义通道（权重 0.5）：HNSW 余弦检索，top_k=100；知识图谱通道（权重 0.2）：实体匹配→邻居遍历，top_k=30。']),
    ('RRF 融合', ['score(chunk) = Σ 1/(k + rank_i)，k=60。三路并行检索 → RRF 加权融合 → 统一排序。']),
    ('文件规划', ['新建 raganything/hybrid_search.py：HybridSearchEngine 类，含三个独立搜索方法和 _rrf_fuse() 方法。']),
    ('验收标准', ['三路独立检索并行执行；RRF 融合后 Hit Rate 相比单通道提升 ≥ 50%（目标值，待实测）；P95 延迟 < 200ms。']),
])

h('3.2  GraphRAG 知识图谱增强检索', 2)
p('✅ 可行性：可行——实体抽取已有（LightRAG processor），图谱存储已有（NetworkX），需做检索管线集成。')
feature_table([
    ('功能目标', ['将现有知识图谱功能从独立可视化升级为深度集成到检索管线，实现基于实体关系的图谱增强检索。']),
    ('图谱检索集成', ['在混合检索第三通道中：识别用户问题中的实体→在知识图谱中匹配→邻居遍历→收集关联文档片段→参与 RRF 融合。']),
    ('前端可视化', ['通过 /api/knowledge/graph 获取图谱数据，使用 D3 力导向布局渲染。支持节点点击展开、拖拽、缩放。']),
    ('验收标准', ['图谱检索通道独立返回结果并参与 RRF 融合；图谱随新文档入库自动更新；前端 D3 力导向图正常渲染。']),
])

h('3.3  DAG 可视化工作流引擎（MVP 版）', 2)
p('🔴 可行性：需降级——完整版（41 种节点 + cron 调度）实际需要 3 周。第 3 周交付 MVP 版：5-8 种核心节点 + 仅手动执行。')
feature_table([
    ('MVP 功能范围', ['节点类型：数据源（文件上传、数据库连接）、清洗（字段筛选、去重）、AI（PDF 解析、文本分段、向量化）、输出（导入知识库）。共 5-8 种核心节点，仅手动触发。']),
    ('前端实现', ['React Flow 拖拽画布：左侧算子面板、右侧节点配置抽屉、中间 DAG 画布（拖拽+贝塞尔连线+环路检测）。页面路径：/admin/workflows。']),
    ('后端实现', ['新建 workflows 和 workflow_runs 两张 PostgreSQL 表。CRUD API：/api/workflows。执行引擎按拓扑顺序依次执行节点，每节点超时 300 秒。']),
    ('完整版（远期）', ['41 种节点类型（含全部 AI 算子）；cron 周期调度（APScheduler）；动态表单配置；任务运行视图与历史回溯。']),
])

h('3.4  SSO / OIDC 企业统一认证', 2)
p('✅ 可行性：可行——authlib 库成熟，OIDC 为标准协议，现有 JWT 认证体系完全兼容。')
feature_table([
    ('功能目标', ['在现有用户名密码登录基础上，增加 OIDC 协议的单点登录支持，兼容 Keycloak、LDAP、OAuth 2.0。']),
    ('技术方案', ['使用 authlib 实现 OIDC Provider。新增 /api/auth/oidc/login（重定向授权）和 /api/auth/oidc/callback（回调交换 Token）。OIDC 用户首次登录自动创建本地用户。']),
    ('环境变量', ['OIDC_ENABLED=true、OIDC_ISSUER、OIDC_CLIENT_ID、OIDC_CLIENT_SECRET。关闭时不影响现有登录方式。']),
    ('验收标准', ['可通过 Keycloak 账号登录；首次 OIDC 登录自动创建本地用户；现有用户名密码登录不受影响。']),
])

h('3.5  多轮对话上下文管理', 2)
p('✅ 可行性：可行——conversations API 已有，滑动窗口+摘要压缩是纯逻辑层改动。')
feature_table([
    ('功能目标', ['为智能体对话增加滑动窗口上下文管理和 LLM 摘要压缩，解决长对话 Token 消耗线性增长的问题。']),
    ('滑动窗口', ['保留最近 10 轮对话（context_window=10）。超出窗口时触发压缩：前 3 轮保留原文，更早的对话压缩为摘要。']),
    ('摘要压缩', ['Token 预算 2000（token_budget=2000）。超出时调用 LLM 将最早历史压缩为摘要，目标压缩比 ≥ 60%。']),
    ('验收标准', ['超过 10 轮对话自动压缩；压缩后 Token 节省 ≥ 60%（目标值）；压缩后仍能正确回答关联问题。']),
])

h('3.6  密钥管理与依赖安全扫描', 2)
p('✅ 可行性：可行——配置+CI 集成类任务，无技术难点。')
feature_table([
    ('密钥管理', ['当前已实现：secrets.token_hex(32) 自动生成密钥、.env 管理、mask_sensitive_data() 日志脱敏。待补充：密钥轮换策略文档（建议每 90 天轮换 JWT_SECRET）、CI 环境变量注入指南。']),
    ('依赖安全扫描', ['CI 方案：新建 .github/workflows/security-scan.yml，每次 PR 运行 pip-audit + npm audit。高危漏洞（CVSS ≥ 7.0）阻断 PR 合并。每周全量扫描并自动创建 Issue。']),
])

h('3.7  核心测试覆盖', 2)
p('✅ 可行性：可行——pytest 已配置，已有 test_auth.py 和 test_core_modules.py 基础测试。')
feature_table([
    ('测试范围', ['认证模块（100% 覆盖率）、知识库 CRUD（80%）、智能问答管线（80%）、智能体管理（80%）、分块策略（90%）。']),
    ('新增文件', ['conftest.py（fixtures）、test_knowledge_api.py、test_upload.py、test_query_pipeline.py、test_agent_manager.py、test_chunking_strategies.py。']),
    ('CI 集成', ['.github/workflows/test.yml：每次 PR 运行 pytest --cov，覆盖率报告上传至 Artifacts。PR 合并前必须通过所有测试。']),
])

doc.add_page_break()
# ═══ 四、第 4 周 ═══
h('四、第 4 周交付项')

h('4.1  审计日志系统', 2)
p('✅ 可行性：可行——SQLite 表 + Starlette 中间件，半天工作量。')
feature_table([
    ('功能目标', ['全操作审计追溯：记录操作人、IP、时间、操作类型、目标资源、结果。支持筛选和 CSV 导出。']),
    ('数据库设计', ['SQLite 表 audit_logs（id, user_id, username, ip_address, action, resource_type, resource_id, detail JSON, status, user_agent, created_at）。索引：user_id、action、created_at。保留期 1 年。']),
    ('中间件', ['Starlette BaseHTTPMiddleware：每个 API 请求处理完成后异步写入审计日志，不阻塞响应。豁免 /api/health。']),
    ('查询导出', ['GET /api/audit-logs（按用户/操作/资源/时间筛选，分页 page_size=50）；GET /api/audit-logs/export（CSV 导出）。']),
])

h('4.2  文档解析管线升级（Docling / Marker 集成）', 2)
p('⚠️ 可行性：可行但需注意依赖冲突——Docling 和 Marker 各自依赖不同版本 PyTorch。建议作为可选依赖。')
feature_table([
    ('功能目标', ['在 MinerU 2.0 基础上增加 Docling（IBM 开源，表格精度高）和 Marker（快速轻量）两种备选解析后端。']),
    ('三引擎对比', ['MinerU 2.0（默认）：版面分析+OCR；Docling：表格识别精度高；Marker：速度快，纯文本批量处理。']),
    ('自动选择', ['PARSER_BACKEND=auto 时：文档包含 >10 个表格→Docling；文档 >50 页纯文本→Marker；其他→MinerU。']),
    ('安装方式', ['可选依赖：pip install raganything[docling] 或 [marker] 或 [all]。默认仅包含 MinerU。']),
])

h('4.3  前端 Zustand + i18n 国际化重构', 2)
p('✅ 可行性：可行——Zustand 和 react-i18next 均为成熟 npm 包。')
feature_table([
    ('状态管理', ['将 React Context 迁移到 Zustand Store：useAuthStore、useKnowledgeStore、useAgentStore、useQueryStore、useUIStore。减少不必要的重渲染。']),
    ('国际化', ['中英文切换：locales/zh-CN/ 和 en-US/ 目录，四个语言包（common/knowledge/agent/query.json）。所有硬编码中文通过 t() 引用。']),
])

h('4.4  性能压测与生产调优', 2)
p('✅ 可行性：可行——locust 为成熟压测工具。')
feature_table([
    ('压测方案', ['使用 locust 编写压测脚本。场景：健康检查（200 并发/5min）、登录（50 并发/5min）、上传（20 并发/10min）、查询（50 并发/10min）、流式问答（10 并发/10min）。']),
    ('调优目标', ['QPS ≥ 50（非 AI 接口）；P95 延迟 < 200ms；流式问答首 Token < 5s。配置：Uvicorn workers=CPU 核数、PostgreSQL pool=20、Redis maxmemory=512MB LRU。']),
    ('产出', ['《性能测试报告》：各场景 QPS/P50/P95/P99 延迟、资源占用、瓶颈分析和优化建议。']),
])

h('4.5  项目管理看板集成', 2)
p('✅ 可行性：可行——独立 HTML 看板已有，移植到 React 组件即可。')
feature_table([
    ('功能目标', ['将独立 HTML 看板集成到主前端应用 /admin/pm 页面，使用 Zustand 管理状态，数据迁移到后端 API。']),
    ('后端 API', ['/api/pm/*：GET /api/pm/tasks、PUT /api/pm/tasks/{id}、GET /api/pm/notes/{id}、PUT /api/pm/notes/{id}。']),
    ('协作', ['通过 WebSocket 实现看板状态实时同步：用户 A 标记完成→推送事件→用户 B 自动更新。']),
])

doc.add_page_break()
# ═══ 五、远期 ═══
h('五、远期规划')
p('以下功能为按需扩展项，启动时间根据实际需求优先级确定。')

h('5.1  数据集模块', 2)
feature_table([
    ('功能范围', ['创建 Clickhouse/Hive 数据集、表结构预览、数据预览（前 1000 条）、SQL 查询（受限 SELECT）、数据血缘追踪、操作记录。']),
    ('技术方案', ['Clickhouse Python 客户端 + SQLAlchemy Hive 方言。SQL 白名单限制（仅 SELECT）。数据血缘通过 sqlparse 解析 AST。']),
])

h('5.2  多数据库接入扩展', 2)
feature_table([
    ('扩展范围', ['在 MySQL + PostgreSQL 基础上增加：Oracle（cx_Oracle）、SQL Server（pyodbc）、TiDB、MongoDB（motor）、ClickHouse、Hive（pyhive）。']),
    ('统一接口', ['BaseDBConnector 抽象基类：connect()、list_tables()、get_schema()、query()、close()。各数据库实现子类。']),
])

h('5.3  知识质量检测', 2)
feature_table([
    ('三项检测', ['错别字检测（pycorrector + LLM 验证）；语句完整性（LTP/HanLP 句法分析）；敏感词检测（本地词库 + AC 自动机多模式匹配）。']),
])

h('5.4  非结构化数据打标应用', 2)
feature_table([
    ('功能范围', ['预置 ASR 转写、图片理解、主体识别、标签打标、情感识别、观点总结、水军识别等 AI 应用。支持自定义打标应用（自定义提示词+输出字段）。标签库管理。错题反馈+提示词自动优化。']),
])

h('5.5  其他远期项目', 2)
feature_table([
    ('更多数据源接入', ['腾讯云文档、钉钉云文档、金山云文档（私有化）、微信素材库。按甲方实际使用工具决定优先级。']),
    ('实时协作', ['多人实时编辑知识库（WebSocket + CRDT）；评论与标注（数据库 + 实时推送）；任务分配（权限模型扩展）。']),
    ('移动端适配', ['微信小程序（Taro/uni-app）；移动 Web 响应式适配。']),
])

# Save
doc.save(r'c:\Users\98014\RAG-Anything\RAG-Anything待开发功能详解.docx')
print('OK')
