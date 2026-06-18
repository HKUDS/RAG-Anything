#!/usr/bin/env python3
"""Convert Markdown files to formatted DOCX documents."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if not text:
        return p
    # Handle inline formatting: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                run = p.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
                run.underline = True
        else:
            run = p.add_run(part)
        if bold: run.bold = True
        if italic: run.italic = True
        if size: run.font.size = size
        if color: run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p

def parse_table_line(line):
    """Parse a markdown table row like | a | b | c |"""
    cells = []
    current = ''
    in_cell = False
    for ch in line:
        if ch == '|':
            if not in_cell:
                in_cell = True
            else:
                cells.append(current.strip())
                current = ''
                in_cell = False
        elif in_cell:
            current += ch
    if current.strip():
        cells.append(current.strip())
    return cells

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Configure styles
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.35

    for level, info in [(1, (22, True)), (2, (16, True)), (3, (13, True)), (4, (11.5, True))]:
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = '微软雅黑'
        hstyle.font.size = Pt(info[0])
        hstyle.font.bold = info[1]
        hstyle.font.color.rgb = RGBColor(0x1A, 0x1D, 0x23)
        if level == 1:
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
        else:
            hstyle.paragraph_format.space_before = Pt(16)
            hstyle.paragraph_format.space_after = Pt(8)

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    i = 0
    in_table = False
    table_rows = []
    table_alignments = []
    in_code_block = False
    code_lines = []
    in_quote = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                for cl in code_lines:
                    p = doc.add_paragraph(style='Normal')
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    run = p.add_run(cl)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x5F, 0x65, 0x70)
            run.italic = True
            i += 1
            # Handle multi-line blockquotes
            while i < len(lines) and lines[i].startswith('> '):
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(lines[i][2:].rstrip())
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x5F, 0x65, 0x70)
                run.italic = True
                i += 1
            continue

        # Horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D0D0"/></w:pBdr>')
            pPr.append(pBdr)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
                table_alignments = []

            cells = parse_table_line(line)

            # Check if it's a separator row like |------|------|
            if all(re.match(r'^:?-{3,}:?$', c) for c in cells if c):
                # Calculate alignments
                table_alignments = []
                for c in cells:
                    c = c.strip()
                    if c.startswith(':') and c.endswith(':'):
                        table_alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                    elif c.endswith(':'):
                        table_alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        table_alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            else:
                table_rows.append(cells)

            i += 1

            # Check if next line continues the table
            if i >= len(lines) or '|' not in lines[i] or not lines[i].strip().startswith('|'):
                # End of table
                in_table = False
                if table_rows:
                    # Create table
                    num_cols = max(len(r) for r in table_rows)
                    num_rows = len(table_rows)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Set column widths (approximate)
                    available_width = Cm(17)
                    col_width = available_width / max(num_cols, 1)

                    for r_idx, row_data in enumerate(table_rows):
                        row = table.rows[r_idx]
                        # Pad to match columns
                        while len(row_data) < num_cols:
                            row_data.append('')
                        for c_idx, cell_text in enumerate(row_data):
                            cell = row.cells[c_idx]
                            cell.width = col_width

                            # Clear default paragraph
                            cell.paragraphs[0].clear()

                            run = cell.paragraphs[0].add_run(cell_text.strip())
                            run.font.name = '微软雅黑'
                            if r_idx == 0:
                                run.font.size = Pt(9)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                set_cell_shading(cell, '1F4E79')
                            else:
                                run.font.size = Pt(8.5)

                            # Alternate row shading
                            if r_idx > 0 and r_idx % 2 == 0:
                                set_cell_shading(cell, 'F7F8FA')

                            # Alignment
                            if r_idx > 0 and c_idx < len(table_alignments):
                                cell.paragraphs[0].alignment = table_alignments[c_idx]

                            # Cell padding
                            tcPr = cell._tc.get_or_add_tcPr()
                            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="40"/><w:left w:w="80"/><w:bottom w:w="40"/><w:right w:w="80"/></w:tcMar>')
                            tcPr.append(tcMar)

                    doc.add_paragraph()  # spacer
                table_rows = []
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Lists
        if re.match(r'^[\s]*[-*]\s', line):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_parts_to_paragraph(p, text)
            # Handle multi-line list items
            i += 1
            while i < len(lines) and lines[i].startswith('  ') and not lines[i].strip().startswith('-') and not lines[i].strip().startswith('*') and '|' not in lines[i]:
                add_parts_to_paragraph(p, ' ' + lines[i].strip())
                i += 1
            continue

        # Numbered lists
        if re.match(r'^[\s]*\d+\.\s', line):
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_parts_to_paragraph(p, text)
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            add_formatted_paragraph(doc, line)
        i += 1

    doc.save(docx_path)
    print(f'OK: {docx_path}')

def add_parts_to_paragraph(p, text):
    """Add inline formatted runs to a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                run = p.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
                run.underline = True
        else:
            p.add_run(part)

if __name__ == '__main__':
    files = [
        (r'c:\Users\98014\RAG-Anything\RAG-Anything已实现功能说明书.md',
         r'c:\Users\98014\RAG-Anything\RAG-Anything已实现功能说明书.docx'),
        (r'c:\Users\98014\RAG-Anything\RAG-Anything待开发功能说明书.md',
         r'c:\Users\98014\RAG-Anything\RAG-Anything待开发功能说明书.docx'),
    ]
    for md_path, docx_path in files:
        try:
            convert_md_to_docx(md_path, docx_path)
        except Exception as e:
            print(f'ERROR converting {md_path}: {e}')
            import traceback
            traceback.print_exc()
